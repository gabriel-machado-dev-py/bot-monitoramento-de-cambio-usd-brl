import platform
from selenium import webdriver
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import NoSuchElementException, TimeoutException
from selenium.webdriver.chrome.options import Options
import logging as lg
from docx2pdf import convert
from datetime import datetime
from docx import Document as DocxDocument
from docx.shared import RGBColor, Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fpdf import FPDF

# Configuração de logging
lg.basicConfig(level=lg.INFO, format='%(asctime)s - %(message)s')

DOWNLOAD_DIRECTORY = 'C:\\Users\\gabri\\OneDrive\\Área de Trabalho\\projetos\\pyautogui\\selenium_dev_aprender'

def iniciar_driver(url_do_site):
    try:
        chrome_options = Options()
        argumentos = [
            '--lang=pt-BR', 
            '--window-size=1300,1000', 
            '--disable-notifications', 
            '--incognito', 
            '--block-new-web-contents', 
            '--no-default-browser-check', 
            'window-position=36,68',
            '--disable-popup-blocking',  
            '--disable-infobars',        
            '--disable-extensions',      
            '--disable-gpu',             
            '--headless'                 
        ]
        for argumento in argumentos:
            chrome_options.add_argument(argumento)

        chrome_options.add_experimental_option('excludeSwitches', ['enable-automation'])
        chrome_options.add_experimental_option('prefs', {
            'download.default_directory': DOWNLOAD_DIRECTORY,
            'download.directory_upgrade': True,
            'download.prompt_for_download': False,
            'profile.default_content_setting_values.notifications': 2,
            'profile.default_content_setting_values.automatic_downloads': 1,
        })

        driver = webdriver.Chrome(options=chrome_options)
        driver.get(url_do_site)

        wait = WebDriverWait(driver, 10, poll_frequency=1, ignored_exceptions=[
            NoSuchElementException,
            TimeoutException
        ])

        return driver, wait
    except Exception as e:
        lg.error(f'Erro ao inicializar o driver: {type(e).__name__} - {e}')
        return None, None

def extrair_cotacao_dolar():
    driver, wait = iniciar_driver('https://valor.globo.com/valor-data/')
    if not driver:
        return None, None, None, None, None

    try:
        cotacao_dolar_element = wait.until(EC.presence_of_element_located(('xpath', "//div[@class='grid-x']")))
        cotacao_dolar_texto = cotacao_dolar_element.text.split("\n")
        url_site = driver.current_url
        driver.save_screenshot('site.png')

        if cotacao_dolar_texto:
            
            cotacao_dolar_titulo = cotacao_dolar_texto[0]
            cotacao_dolar_valor = float(cotacao_dolar_texto[1].replace(',', '.'))
            cotacao_dolar_porcentagem = float(cotacao_dolar_texto[2].replace('-', '').replace(',', '.').replace('%', ''))
            
            data_atual = datetime.now().strftime('%d/%m/%Y')
            
            lg.info(f'Cotação do dólar: {cotacao_dolar_valor:.2f} - {cotacao_dolar_titulo} - {cotacao_dolar_porcentagem}% - {data_atual}')
            
            return cotacao_dolar_valor, cotacao_dolar_titulo, cotacao_dolar_porcentagem, data_atual, url_site
    except NoSuchElementException as e:
        lg.error(f'Elemento não encontrado: {e}')
    except TimeoutException as e:
        lg.error(f'Tempo de espera excedido: {e}')
    except Exception as e:
        lg.error(f'Erro ao extrair cotação do dólar: {type(e).__name__} - {e}')

    return None, None, None, None, None

def salvar_dados_word():
    cotacao_dolar_valor, cotacao_dolar_titulo, cotacao_dolar_porcentagem, data_atual, url_site = extrair_cotacao_dolar()
    if not cotacao_dolar_valor:
        return

    doc = DocxDocument()

    # Título
    heading_text = f'{cotacao_dolar_titulo} - R${cotacao_dolar_valor:.2f} ({data_atual})\n\n'
    heading = doc.add_heading(heading_text, level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # Verde
        run.font.size = Pt(24)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Parágrafo 1
    paragraph1 = doc.add_paragraph()
    run1 = paragraph1.add_run(f'O dólar está custando R${cotacao_dolar_valor:.2f} na data de {data_atual}.')
    run1.font.size = Pt(12)
    paragraph1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph1.paragraph_format.space_after = Pt(0)

    # Parágrafo 2
    paragraph2 = doc.add_paragraph()
    run2 = paragraph2.add_run(f'Valor cotado no site: {url_site}')
    run2.font.size = Pt(12)
    paragraph2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Imagem
    doc.add_picture('site.png', width=Cm(15))

    # Parágrafo 3
    if cotacao_dolar_porcentagem > 0:
        paragraph3 = doc.add_paragraph(f'Cotação caiu {cotacao_dolar_porcentagem}% em relação ao dia anterior.\n\n')
    elif cotacao_dolar_porcentagem < 0:
        paragraph3 = doc.add_paragraph(f'Cotação subiu {cotacao_dolar_porcentagem}% em relação ao dia anterior.\n\n')
    else:
        paragraph3 = doc.add_paragraph('Cotação está estável em relação ao dia anterior.\n\n')
    paragraph3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Parágrafo 4
    paragraph4 = doc.add_paragraph()
    run4 = paragraph4.add_run('Cotação feita por: Gabriel Machado')
    run4.font.color.rgb = RGBColor(0x80, 0x00, 0x80)  # Roxo
    run4.font.size = Pt(12)
    paragraph4.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save('cotacao_dolar.docx')

def converter_word_pdf():
    sistema_operacional = platform.system()
    
    if sistema_operacional == 'Windows':
        convert("cotacao_dolar.docx", "cotacao_dolar.pdf")
        
    elif sistema_operacional in ['Darwin', 'Linux']:
        # Usar python-docx e fpdf2 para macOS e Linux
        doc = DocxDocument("cotacao_dolar.docx")
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        
        for para in doc.paragraphs:
            pdf.multi_cell(0, 10, para.text)
        
        pdf.output("cotacao_dolar.pdf")
    else:
        print("Sistema operacional não suportado.")

if __name__ == '__main__':
    salvar_dados_word()
    converter_word_pdf()
    